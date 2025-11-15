# -*- coding: utf-8 -*-
"""
Streamlit app: Thẩm định phương án kinh doanh/ sử dụng vốn (pasdv.docx)
"""
import io
import os
import re
import math
import json
import zipfile
import datetime as dt
from typing import Dict, Any, Tuple, Optional
import numpy as np
import pandas as pd
import streamlit as st
# Docx parsing
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except Exception:
    Document = None
# Gemini
try:
    import google.generativeai as genai
except Exception:
    genai = None
# Plotly cho biểu đồ
try:
    import plotly.graph_objects as go
    import plotly.express as px
except Exception:
    go = None
    px = None
st.set_page_config(page_title="PASDV - Thẩm định phương án", page_icon="💼", layout="wide")
# ========================== Helpers ==========================
FIELD_DEFAULTS = {
    "ten_khach_hang": "",
    "cccd": "",
    "noi_cu_tru": "",
    "so_dien_thoai": "",
    "muc_dich_vay": "",
    "tong_nhu_cau_von": 0.0,
    "von_doi_ung": 0.0,
    "so_tien_vay": 0.0,
    "lai_suat_nam": 10.0,
    "thoi_gian_vay_thang": 12,
    "ky_han_tra": "Tháng",
    "thu_nhap_thang": 0.0,
    "gia_tri_tsdb": 0.0,
    "tong_no_hien_tai": 0.0,
    "loi_nhuan_rong_nam": 0.0,
    "tong_von_dau_tu": 0.0,
}
def vnd_to_float(s: str) -> float:
    """Chuyển chuỗi tiền tệ VN về float (hỗ trợ dấu . ngăn cách, , thập phân)."""
    if s is None:
        return 0.0
    s = str(s)
    if "," in s and "." in s:
        s = s.replace(".", "").replace(",", ".")
    elif "," in s and "." not in s:
        s = s.replace(".", "")
        s = s.replace(",", ".")
    else:
        s = s.replace(".", "")
    s = s.replace("đ", "").replace("VND", "").replace("vnđ", "").replace("₫", "").replace(" ", "")
    s = re.sub(r"[^\d\.\-]", "", s)
    try:
        return float(s) if s else 0.0
    except Exception:
        return 0.0
def format_vnd(amount: float) -> str:
    """Định dạng tiền VND: 1.234.567"""
    try:
        return f"{float(amount):,.0f}".replace(",", ".")
    except Exception:
        return "0"
def format_vnd_float(amount: float) -> str:
    """Định dạng số thập phân kiểu VN: 1.234.567,89"""
    try:
        s = f"{float(amount):,.2f}"
        s = s.replace(",", "_").replace(".", ",").replace("_", ".")
        return s
    except Exception:
        return "0,00"
def percent_to_float(s: str) -> float:
    """Chuyển đổi chuỗi phần trăm sang số float; chấp nhận '8,5' hoặc '8.5'."""
    if s is None:
        return 0.0
    s = str(s).replace(",", ".")
    m = re.search(r"(\d+(?:\.\d+)?)", s)
    return float(m.group(1)) if m else 0.0
def vn_money_input(label: str, value: float, key: Optional[str] = None, help: Optional[str] = None) -> float:
    """Ô nhập tiền tệ kiểu VN: hiển thị 1.234.567 và parse lại về float."""
    raw = st.text_input(label, value=format_vnd(value), key=key, help=help)
    return float(vnd_to_float(raw))
def vn_percent_input(label: str, value: float, key: Optional[str] = None, help: Optional[str] = None) -> float:
    """Ô nhập phần trăm linh hoạt: cho phép nhập '8,5' hoặc '8.5'."""
    shown = f"{float(value):.2f}".replace(".", ",")
    raw = st.text_input(label, value=shown, key=key, help=help)
    return percent_to_float(raw)
def extract_from_docx(file_bytes: bytes) -> Dict[str, Any]:
    """Đọc .docx PASDV và trích xuất thông tin theo cấu trúc thực tế."""
    data = FIELD_DEFAULTS.copy()
    if Document is None:
        return data
    bio = io.BytesIO(file_bytes)
    doc = Document(bio)
    full_text = "\n".join([p.text for p in doc.paragraphs])
    lines = [line.strip() for line in full_text.split('\n') if line.strip()]
    full_text = "\n".join(lines)
    # === 1. THÔNG TIN KHÁCH HÀNG ===
    ten_pattern1 = r"(?:\d+\.\s*)?Họ\s+và\s+tên\s*[:：]\s*([A-ZÀÁẢÃẠĂẰẮẲẴẶÂẦẤẨẪẬĐÈÉẺẼẸÊỀẾỂỄỆÌÍỈĨỊÒÓỎÕỌÔỒỐỔỖỘƠỜỚỞỠỢÙÚỦŨỤƯỪỨỬỮỰỲÝỶỸỴ][a-zàáảãạăằắẳẵặâầấẨẫậđèéẻẽẹêềếểễệìíỉĩịòóỏõọôồốổỗộơờớởỡợùúủũụưừứửữựỳýỷỹỵA-ZÀÁẢÃẠĂẰẮẲẴẶÂẦẤẨẪẬĐÈÉẺẼẸÊỀẾỂỄỆÌÍỈĨỊÒÓỎÕỌÔỒỐỔỖỘƠỜỚỞỠỢÙÚỦŨỤƯỪỨỬỮỰỲÝỶỸỴ\s]+)"
    m = re.search(ten_pattern1, full_text, flags=re.IGNORECASE)
    if m:
        data["ten_khach_hang"] = m.group(1).strip()
    else:
        ten_pattern2 = r"(?:Ông|Bà)\s*$$ (?:bà|ông) $$\s*[:：]\s*([A-ZÀÁẢÃẠĂẰẮẲẴẶÂẦẤẨẪẬĐÈÉẺẼẸÊỀẾỂỄỆÌÍỈĨỊÒÓỎÕỌÔỒỐỔỖỘƠỜỚỞỠỢÙÚỦŨỤƯỪỨỬỮỰỲÝỶỸỴ][a-zàáảãạăằắẳẵặâầấẨẫậđèéẻẽẹêềếểễệìíỉĩịòóỏõọôồốổỗộơờớởỡợùúủũụưừứửữựỳýỷỹỵA-ZÀÁẢÃẠĂẰẮẲẴẶÂẦẤẨẪẬĐÈÉẺẼẸÊỀẾỂỄỆÌÍỈĨỊÒÓỎÕỌÔỒỐỔỖỘƠỜỚỞỠỢÙÚỦŨỤƯỪỨỬỮỰỲÝỶỸỴ\s]+)"
        m = re.search(ten_pattern2, full_text, flags=re.IGNORECASE)
        if m:
            data["ten_khach_hang"] = m.group(1).strip()
   
    cccd_pattern = r"(?:CMND|CCCD)(?:\/(?:CCCD|CMND))?(?:\/hộ\s*chiếu)?\s*[:：]\s*(\d{9,12})"
    m = re.search(cccd_pattern, full_text, flags=re.IGNORECASE)
    if m:
        data["cccd"] = m.group(1).strip()
    noi_cu_tru_pattern = r"Nơi\s*cư\s*trú\s*[:：]\s*([^\n]+?)(?=\n|Số\s*điện\s*thoại|$)"
    m = re.search(noi_cu_tru_pattern, full_text, flags=re.IGNORECASE | re.DOTALL)
    if m:
        data["noi_cu_tru"] = m.group(1).strip()
    sdt_pattern = r"Số\s*điện\s*thoại\s*[:：]\s*(0\d{9,10})"
    m = re.search(sdt_pattern, full_text, flags=re.IGNORECASE)
    if m:
        data["so_dien_thoai"] = m.group(1).strip()
    # === 2. PHƯƠNG ÁN SỬ DỤNG VỐN ===
    muc_dich_pattern1 = r"Mục\s*đích\s*vay\s*[:：]\s*([^\n]+)"
    m = re.search(muc_dich_pattern1, full_text, flags=re.IGNORECASE)
    if m:
        data["muc_dich_vay"] = m.group(1).strip()
    else:
        muc_dich_pattern2 = r"Vốn\s*vay\s*Agribank.*?[:：].*?(?:Thực\s*hiện|Sử\s*dụng\s*vào)\s*([^\n]+)"
        m = re.search(muc_dich_pattern2, full_text, flags=re.IGNORECASE | re.DOTALL)
        if m:
            data["muc_dich_vay"] = m.group(1).strip()[:200]
    tnc_pattern = r"(?:Tổng\s*nhu\s*cầu\s*vốn|1\.\s*Tổng\s*nhu\s*cầu\s*vốn)\s*[:：]\s*([\d\.,]+)"
    m = re.search(tnc_pattern, full_text, flags=re.IGNORECASE)
    if m:
        data["tong_nhu_cau_von"] = vnd_to_float(m.group(1))
    von_du_pattern = r"Vốn\s*đối\s*ứng\s*(?:tham\s*gia)?[^\d]*([\d\.,]+)\s*đồng"
    m = re.search(von_du_pattern, full_text, flags=re.IGNORECASE)
    if m:
        data["von_doi_ung"] = vnd_to_float(m.group(1))
    so_tien_vay_pattern = r"Vốn\s*vay\s*Agribank\s*(?:số\s*tiền)?[:\s]*([\d\.,]+)\s*đồng"
    m = re.search(so_tien_vay_pattern, full_text, flags=re.IGNORECASE)
    if m:
        data["so_tien_vay"] = vnd_to_float(m.group(1))
    thoi_han_pattern = r"Thời\s*hạn\s*vay\s*[:：]\s*(\d+)\s*tháng"
    m = re.search(thoi_han_pattern, full_text, flags=re.IGNORECASE)
    if m:
        data["thoi_gian_vay_thang"] = int(m.group(1))
    lai_suat_pattern = r"Lãi\s*suất\s*[:：]\s*([\d\.,]+)\s*%"
    m = re.search(lai_suat_pattern, full_text, flags=re.IGNORECASE)
    if m:
        data["lai_suat_nam"] = percent_to_float(m.group(1))
    # === 3. NGUỒN TRẢ NỢ & THU NHẬP ===
    thu_nhap_du_an_pattern = r"Từ\s*nguồn\s*thu\s*của\s*dự\s*án[^\d]*([\d\.,]+)\s*đồng\s*/\s*tháng"
    m = re.search(thu_nhap_du_an_pattern, full_text, flags=re.IGNORECASE)
    thu_nhap_du_an = 0.0
    if m:
        thu_nhap_du_an = vnd_to_float(m.group(1))
    thu_nhap_luong_pattern = r"Thu\s*nhập\s*từ\s*lương\s*[:：]\s*([\d\.,]+)\s*đồng\s*/\s*tháng"
    m = re.search(thu_nhap_luong_pattern, full_text, flags=re.IGNORECASE)
    thu_nhap_luong = 0.0
    if m:
        thu_nhap_luong = vnd_to_float(m.group(1))
    tong_thu_nhap_pattern = r"Tổng\s*thu\s*nhập\s*(?:ổn\s*định)?\s*(?:hàng\s*)?tháng\s*[:：]\s*([\d\.,]+)\s*đồng"
    m = re.search(tong_thu_nhap_pattern, full_text, flags=re.IGNORECASE)
    if m:
        data["thu_nhap_thang"] = vnd_to_float(m.group(1))
    else:
        data["thu_nhap_thang"] = thu_nhap_luong + thu_nhap_du_an
    # === 4. TÀI SẢN BẢO ĐẢM ===
    tsdb_pattern1 = r"Tài\s*sản\s*1[^\n]*Giá\s*trị\s*[:：]\s*([\d\.,]+)\s*đồng"
    m = re.search(tsdb_pattern1, full_text, flags=re.IGNORECASE | re.DOTALL)
    if m:
        data["gia_tri_tsdb"] = vnd_to_float(m.group(1))
    else:
        tsdb_pattern2 = r"Giá\s*trị\s*nhà\s*dự\s*kiến\s*mua\s*[:：]\s*([\d\.,]+)\s*đồng"
        m = re.search(tsdb_pattern2, full_text, flags=re.IGNORECASE)
        if m:
            data["gia_tri_tsdb"] = vnd_to_float(m.group(1))
    # === 5. THÔNG TIN BỔ SUNG ===
    loi_nhuan_pattern = r"Lợi\s*nhuận\s*(?:ròng)?\s*(?:năm)?[^\d]*([\d\.,]+)\s*đồng"
    m = re.search(loi_nhuan_pattern, full_text, flags=re.IGNORECASE)
    if m:
        data["loi_nhuan_rong_nam"] = vnd_to_float(m.group(1))
    elif thu_nhap_du_an > 0:
        data["loi_nhuan_rong_nam"] = thu_nhap_du_an * 12
    if data["tong_nhu_cau_von"] == 0 and (data["von_doi_ung"] + data["so_tien_vay"] > 0):
        data["tong_nhu_cau_von"] = data["von_doi_ung"] + data["so_tien_vay"]
    if data["tong_von_dau_tu"] == 0:
        data["tong_von_dau_tu"] = data["tong_nhu_cau_von"]
    if data["gia_tri_tsdb"] == 0 and data["tong_nhu_cau_von"] > 0:
        data["gia_tri_tsdb"] = data["tong_nhu_cau_von"]
    return data
def annuity_payment(principal: float, annual_rate_pct: float, months: int) -> float:
    r = annual_rate_pct / 100.0 / 12.0
    if months <= 0:
        return 0.0
    if r == 0:
        return principal / months
    pmt = principal * r * (1 + r) ** months / ((1 + r) ** months - 1)
    return pmt
def build_amortization(principal: float, annual_rate_pct: float, months: int, start_date: Optional[dt.date] = None) -> pd.DataFrame:
    if start_date is None:
        start_date = dt.date.today()
    r = annual_rate_pct / 100.0 / 12.0
    pmt = annuity_payment(principal, annual_rate_pct, months)
    schedule = []
    balance = principal
    for i in range(1, months + 1):
        interest = balance * r
        principal_pay = pmt - interest
        balance = max(0.0, balance - principal_pay)
        pay_date = start_date + dt.timedelta(days=30 * i)
        schedule.append({
            "Kỳ": i,
            "Ngày thanh toán": pay_date.strftime("%d/%m/%Y"),
            "Tiền lãi": round(interest, 0),
            "Tiền gốc": round(principal_pay, 0),
            "Tổng phải trả": round(pmt, 0),
            "Dư nợ còn lại": round(balance, 0),
        })
    df = pd.DataFrame(schedule)
    return df
def style_schedule_table(df: pd.DataFrame) -> pd.DataFrame:
    """Tô màu bảng kế hoạch trả nợ"""
    def color_row(row):
        if row['Kỳ'] % 2 == 0:
            return ['background-color: #f0f8ff'] * len(row)
        else:
            return ['background-color: #ffffff'] * len(row)
    styled = df.style.apply(color_row, axis=1)
    styled = styled.format({
        'Tiền lãi': lambda x: format_vnd(x),
        'Tiền gốc': lambda x: format_vnd(x),
        'Tổng phải trả': lambda x: format_vnd(x),
        'Dư nợ còn lại': lambda x: format_vnd(x)
    })
    styled = styled.set_properties(**{
        'text-align': 'right',
        'font-size': '14px'
    }, subset=['Tiền lãi', 'Tiền gốc', 'Tổng phải trả', 'Dư nợ còn lại'])
    styled = styled.set_properties(**{
        'text-align': 'center'
    }, subset=['Kỳ', 'Ngày thanh toán'])
    return styled
def compute_metrics(d: Dict[str, Any]) -> Dict[str, Any]:
    res = {}
    pmt = annuity_payment(d.get("so_tien_vay", 0.0), d.get("lai_suat_nam", 0.0), d.get("thoi_gian_vay_thang", 0))
    thu_nhap_thang = max(1e-9, d.get("thu_nhap_thang", 0.0))
    res["PMT_thang"] = pmt
    res["DSR"] = pmt / thu_nhap_thang if thu_nhap_thang > 0 else np.nan
    tong_nhu_cau = d.get("tong_nhu_cau_von", 0.0)
    von_doi_ung = d.get("von_doi_ung", 0.0)
    so_tien_vay = d.get("so_tien_vay", 0.0)
    gia_tri_tsdb = d.get("gia_tri_tsdb", 0.0)
    tong_no_hien_tai = d.get("tong_no_hien_tai", 0.0)
    loi_nhuan_rong_nam = d.get("loi_nhuan_rong_nam", 0.0)
    tong_von_dau_tu = d.get("tong_von_dau_tu", 0.0)
    res["E_over_C"] = (von_doi_ung / tong_nhu_cau) if tong_nhu_cau > 0 else np.nan
    res["LTV"] = (so_tien_vay / gia_tri_tsdb) if gia_tri_tsdb > 0 else np.nan
    thu_nhap_nam = thu_nhap_thang * 12.0
    res["Debt_over_Income"] = (tong_no_hien_tai + so_tien_vay) / max(1e-9, thu_nhap_nam)
    res["ROI"] = (loi_nhuan_rong_nam / max(1e-9, tong_von_dau_tu)) if tong_von_dau_tu > 0 else np.nan
    res["CFR"] = ((thu_nhap_thang - pmt) / thu_nhap_thang) if thu_nhap_thang > 0 else np.nan
    res["Coverage"] = (gia_tri_tsdb / max(1e-9, so_tien_vay)) if so_tien_vay > 0 else np.nan
    res["Phuong_an_hop_ly"] = math.isclose(tong_nhu_cau, von_doi_ung + so_tien_vay, rel_tol=0.02, abs_tol=1_000_000)
    score = 0.0
    if not np.isnan(res["DSR"]):
        score += max(0.0, 1.0 - min(1.0, res["DSR"])) * 0.25
    if not np.isnan(res["LTV"]):
        score += max(0.0, 1.0 - min(1.0, res["LTV"])) * 0.25
    if not np.isnan(res["E_over_C"]):
        score += min(1.0, res["E_over_C"] / 0.3) * 0.2
    if not np.isnan(res["CFR"]):
        score += max(0.0, min(1.0, (res["CFR"]))) * 0.2
    if not np.isnan(res["Coverage"]):
        score += min(1.0, (res["Coverage"] / 1.5)) * 0.1
    res["Score_AI_demo"] = round(score, 3)
    return res
def create_metrics_chart(metrics: Dict[str, Any]):
    """Tạo biểu đồ trực quan cho các chỉ tiêu tài chính chính"""
    if go is None or px is None:
        st.warning("Thư viện Plotly chưa được cài đặt. Không thể vẽ biểu đồ.")
        return
    df_metrics = pd.DataFrame({
        "Chỉ tiêu": ["DSR", "LTV", "E/C", "Coverage", "CFR"],
        "Giá trị": [
            metrics.get("DSR", np.nan),
            metrics.get("LTV", np.nan),
            metrics.get("E_over_C", np.nan),
            metrics.get("Coverage", np.nan),
            metrics.get("CFR", np.nan),
        ],
        "Ngưỡng tham chiếu": [0.8, 0.8, 0.2, 1.2, 0.0]
    })
    df_metrics = df_metrics.dropna(subset=['Giá trị']).reset_index(drop=True)
    if df_metrics.empty:
        st.info("Không có đủ dữ liệu để vẽ biểu đồ chỉ tiêu tài chính.")
        return
    def get_color(row):
        metric = row['Chỉ tiêu']
        value = row['Giá trị']
        ref = row['Ngưỡng tham chiếu']
        if metric in ["DSR", "LTV"]:
            return "green" if value <= ref else "red"
        elif metric in ["E/C", "Coverage", "CFR"]:
            return "green" if value >= ref else "red"
        return "gray"
    df_metrics['Màu'] = df_metrics.apply(get_color, axis=1)
    df_metrics['Giá trị (%)'] = df_metrics['Giá trị'] * 100
    fig = px.bar(
        df_metrics,
        x="Chỉ tiêu",
        y="Giá trị (%)",
        color="Màu",
        color_discrete_map={"green": "#28a745", "red": "#dc3545", "gray": "#6c757d"},
        text=df_metrics['Giá trị (%)'].apply(lambda x: f"{x:,.1f}%"),
        title="Biểu đồ Chỉ tiêu Tài chính (CADAP)",
        labels={"Giá trị (%)": "Giá trị (%)", "Chỉ tiêu": "Chỉ tiêu"},
    )
    for index, row in df_metrics.iterrows():
        metric = row['Chỉ tiêu']
        ref_value = row['Ngưỡng tham chiếu'] * 100
        color = "#ffc107" if ref_value > 0 else "#007bff"
        if metric in ["DSR", "LTV"]:
            fig.add_shape(
                type="line",
                x0=index - 0.4, x1=index + 0.4, y0=ref_value, y1=ref_value,
                line=dict(color=color, width=2, dash="dash"),
                xref="x", yref="y",
                name=f"Ngưỡng {metric}"
            )
            fig.add_annotation(
                x=index, y=ref_value + 3,
                text=f"Max {ref_value:g}%", showarrow=False,
                font=dict(color=color, size=10),
            )
        elif metric in ["E/C", "Coverage"]:
            fig.add_shape(
                type="line",
                x0=index - 0.4, x1=index + 0.4, y0=ref_value, y1=ref_value,
                line=dict(color=color, width=2, dash="dash"),
                xref="x", yref="y",
                name=f"Ngưỡng {metric}"
            )
            fig.add_annotation(
                x=index, y=ref_value - 3,
                text=f"Min {ref_value:g}%", showarrow=False,
                font=dict(color=color, size=10),
            )
    fig.update_layout(
        showlegend=False,
        yaxis_title="Giá trị (%)",
        xaxis_title="Chỉ tiêu",
        hovermode="x unified"
    )
    st.plotly_chart(fig, use_container_width=True)
def gemini_analyze(d: Dict[str, Any], metrics: Dict[str, Any], model_name: str, api_key: str) -> str:
    if genai is None:
        return "Thư viện google-generativeai chưa được cài. Vui lòng thêm 'google-generativeai' vào requirements.txt."
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        d_formatted = {k: format_vnd(v) if isinstance(v, (int, float)) and k != 'lai_suat_nam' else v for k, v in d.items()}
        metrics_formatted = {
            k: (f"{v*100:,.1f}%"
                if k not in ["PMT_thang", "Debt_over_Income", "Score_AI_demo"] and not np.isnan(v)
                else format_vnd(v) if k == "PMT_thang"
                else f"{v:,.2f}")
            for k, v in metrics.items()
        }
        prompt = f"""
Bạn là chuyên viên tín dụng. Phân tích hồ sơ vay sau (JSON) và đưa ra đề xuất "Cho vay" / "Cho vay có điều kiện" / "Không cho vay" kèm giải thích ngắn gọn (<=200 từ).
JSON đầu vào:
Khách hàng & phương án: {json.dumps(d_formatted, ensure_ascii=False)}
Chỉ tiêu tính toán: {json.dumps(metrics_formatted, ensure_ascii=False)}
Ngưỡng tham chiếu:
- DSR ≤ 0.8; LTV ≤ 0.8; E/C ≥ 0.2; CFR > 0; Coverage > 1.2.
- Nếu thông tin thiếu, hãy nêu giả định rõ ràng.
"""
        resp = model.generate_content(prompt)
        return resp.text or "(Không có nội dung từ Gemini)"
    except Exception as e:
        return f"Lỗi khi gọi Gemini: {e}"
def make_zip_for_download() -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        for fname in ["app.py", "requirements.txt", "README.md"]:
            if os.path.exists(fname):
                z.write(fname, arcname=fname)
    buf.seek(0)
    return buf.read()
def export_to_docx(data: Dict[str, Any], metrics: Dict[str, Any], schedule_df: pd.DataFrame, analysis: str = "") -> bytes:
    """Xuất báo cáo thẩm định ra file DOCX với định dạng đẹp mắt"""
    if Document is None:
        return b""
    
    doc = Document()
    
    # Thiết lập font mặc định
    from docx.oxml.ns import qn
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    doc.styles['Normal'].font.size = Pt(12)
    
    # Trang bìa
    doc.add_paragraph("NGÂN HÀNG NÔNG NGHIỆP VÀ PHÁT TRIỂN NÔNG THÔN VIỆT NAM", style='Title').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("BÁO CÁO THẨM ĐỊNH PHƯƠNG ÁN SỬ DỤNG VỐN", style='Title').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Ngày báo cáo: {dt.date.today().strftime('%d/%m/%Y')}", style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_page_break()
    
    # PHẦN I: THÔNG TIN KHÁCH HÀNG
    doc.add_heading('I. THÔNG TIN KHÁCH HÀNG', 1)
    table1 = doc.add_table(rows=5, cols=2)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.autofit = True
    for row in table1.rows:
        for cell in row.cells:
            cell.paragraphs[0].style.font.size = Pt(11)
            cell.paragraphs[0].style.font.name = 'Times New Roman'
    
    cells = table1.rows[0].cells
    cells[0].text = 'Họ và tên:'
    cells[1].text = data.get('ten_khach_hang', '')
    
    cells = table1.rows[1].cells
    cells[0].text = 'CMND/CCCD:'
    cells[1].text = data.get('cccd', '')
    
    cells = table1.rows[2].cells
    cells[0].text = 'Nơi cư trú:'
    cells[1].text = data.get('noi_cu_tru', '')
    
    cells = table1.rows[3].cells
    cells[0].text = 'Số điện thoại:'
    cells[1].text = data.get('so_dien_thoai', '')
    
    cells = table1.rows[4].cells
    cells[0].text = 'Mục đích vay:'
    cells[1].text = data.get('muc_dich_vay', '')
    
    doc.add_paragraph()
    
    # PHẦN II: THÔNG TIN KHOẢN VAY
    doc.add_heading('II. THÔNG TIN KHOẢN VAY', 1)
    table2 = doc.add_table(rows=7, cols=2)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.autofit = True
    for row in table2.rows:
        for cell in row.cells:
            cell.paragraphs[0].style.font.size = Pt(11)
            cell.paragraphs[0].style.font.name = 'Times New Roman'
    
    cells = table2.rows[0].cells
    cells[0].text = 'Tổng nhu cầu vốn:'
    cells[1].text = f"{format_vnd(data.get('tong_nhu_cau_von', 0))} VND"
    
    cells = table2.rows[1].cells
    cells[0].text = 'Vốn đối ứng:'
    cells[1].text = f"{format_vnd(data.get('von_doi_ung', 0))} VND"
    
    cells = table2.rows[2].cells
    cells[0].text = 'Số tiền vay:'
    cells[1].text = f"{format_vnd(data.get('so_tien_vay', 0))} VND"
    
    cells = table2.rows[3].cells
    cells[0].text = 'Lãi suất:'
    cells[1].text = f"{data.get('lai_suat_nam', 0):.2f}%/năm"
    
    cells = table2.rows[4].cells
    cells[0].text = 'Thời hạn vay:'
    cells[1].text = f"{data.get('thoi_gian_vay_thang', 0)} tháng"
    
    cells = table2.rows[5].cells
    cells[0].text = 'Thu nhập tháng:'
    cells[1].text = f"{format_vnd(data.get('thu_nhap_thang', 0))} VND"
    
    cells = table2.rows[6].cells
    cells[0].text = 'Giá trị TSĐB:'
    cells[1].text = f"{format_vnd(data.get('gia_tri_tsdb', 0))} VND"
    
    doc.add_paragraph()
    
    # PHẦN III: CHỈ TIÊU TÀI CHÍNH
    doc.add_heading('III. CHỈ TIÊU TÀI CHÍNH (CADAP)', 1)
    table3 = doc.add_table(rows=8, cols=3)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    table3.autofit = True
    for i, row in enumerate(table3.rows):
        for cell in row.cells:
            cell.paragraphs[0].style.font.size = Pt(11)
            cell.paragraphs[0].style.font.name = 'Times New Roman'
            if i == 0:  # Header row
                cell.paragraphs[0].style.font.bold = True
                cell.paragraphs[0].style.font.color.rgb = RGBColor(0, 0, 0)
            elif i % 2 == 0:
                cell.paragraphs[0].style.font.color.rgb = RGBColor(0, 0, 139)
    
    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = 'Chỉ tiêu'
    hdr_cells[1].text = 'Giá trị'
    hdr_cells[2].text = 'Đánh giá'
    
    cells = table3.rows[1].cells
    cells[0].text = 'PMT (Tiền trả/tháng)'
    cells[1].text = f"{format_vnd(metrics.get('PMT_thang', 0))} VND"
    cells[2].text = ''
    
    cells = table3.rows[2].cells
    cells[0].text = 'DSR (Debt Service Ratio)'
    dsr = metrics.get('DSR', 0)
    cells[1].text = f"{dsr*100:.1f}%" if not np.isnan(dsr) else 'n/a'
    cells[2].text = '✓ Đạt' if (not np.isnan(dsr) and dsr <= 0.8) else '✗ Không đạt'
    
    cells = table3.rows[3].cells
    cells[0].text = 'LTV (Loan to Value)'
    ltv = metrics.get('LTV', 0)
    cells[1].text = f"{ltv*100:.1f}%" if not np.isnan(ltv) else 'n/a'
    cells[2].text = '✓ Đạt' if (not np.isnan(ltv) and ltv <= 0.8) else '✗ Không đạt'
    
    cells = table3.rows[4].cells
    cells[0].text = 'E/C (Equity to Capital)'
    ec = metrics.get('E_over_C', 0)
    cells[1].text = f"{ec*100:.1f}%" if not np.isnan(ec) else 'n/a'
    cells[2].text = '✓ Đạt' if (not np.isnan(ec) and ec >= 0.2) else '✗ Không đạt'
    
    cells = table3.rows[5].cells
    cells[0].text = 'CFR (Cash Flow Ratio)'
    cfr = metrics.get('CFR', 0)
    cells[1].text = f"{cfr*100:.1f}%" if not np.isnan(cfr) else 'n/a'
    cells[2].text = '✓ Đạt' if (not np.isnan(cfr) and cfr > 0) else '✗ Không đạt'
    
    cells = table3.rows[6].cells
    cells[0].text = 'Coverage (Collateral Coverage)'
    cov = metrics.get('Coverage', 0)
    cells[1].text = f"{cov*100:.1f}%" if not np.isnan(cov) else 'n/a'
    cells[2].text = '✓ Đạt' if (not np.isnan(cov) and cov > 1.2) else '✗ Không đạt'
    
    cells = table3.rows[7].cells
    cells[0].text = 'Score tổng hợp'
    cells[1].text = f"{metrics.get('Score_AI_demo', 0):.3f}"
    score = metrics.get('Score_AI_demo', 0)
    cells[2].text = '✓ Tốt' if score >= 0.7 else ('⚠ Trung bình' if score >= 0.5 else '✗ Yếu')
    
    doc.add_paragraph()
    
    # PHẦN IV: KẾ HOẠCH TRẢ NỢ (5 kỳ đầu)
    doc.add_heading('IV. KẾ HOẠCH TRẢ NỢ (5 kỳ đầu)', 1)
    n_rows = min(6, len(schedule_df) + 1)
    table4 = doc.add_table(rows=n_rows, cols=6)
    table4.style = 'Table Grid'
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    table4.autofit = True
    for i, row in enumerate(table4.rows):
        for cell in row.cells:
            cell.paragraphs[0].style.font.size = Pt(11)
            cell.paragraphs[0].style.font.name = 'Times New Roman'
            if i == 0:  # Header row
                cell.paragraphs[0].style.font.bold = True
            elif i % 2 == 0:
                cell.paragraphs[0].style.font.color.rgb = RGBColor(0, 0, 139)
    
    hdr_cells = table4.rows[0].cells
    hdr_cells[0].text = 'Kỳ'
    hdr_cells[1].text = 'Ngày'
    hdr_cells[2].text = 'Tiền lãi'
    hdr_cells[3].text = 'Tiền gốc'
    hdr_cells[4].text = 'Tổng trả'
    hdr_cells[5].text = 'Dư nợ'
    
    for i in range(min(5, len(schedule_df))):
        row = schedule_df.iloc[i]
        cells = table4.rows[i+1].cells
        cells[0].text = str(row['Kỳ'])
        cells[1].text = row['Ngày thanh toán']
        cells[2].text = format_vnd(row['Tiền lãi'])
        cells[3].text = format_vnd(row['Tiền gốc'])
        cells[4].text = format_vnd(row['Tổng phải trả'])
        cells[5].text = format_vnd(row['Dư nợ còn lại'])
    
    doc.add_paragraph()
    p = doc.add_paragraph(f"(Xem file Excel đính kèm để có đầy đủ {len(schedule_df)} kỳ thanh toán)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # PHẦN V: PHÂN TÍCH VÀ KẾT LUẬN
    if analysis:
        doc.add_heading('V. PHÂN TÍCH VÀ KẾT LUẬN (AI)', 1)
        p = doc.add_paragraph(analysis)
        p.style.font.size = Pt(11)
        doc.add_paragraph()
    
    # PHẦN VI: Ý KIẾN THẨM ĐỊNH
    doc.add_heading('VI. Ý KIẾN THẨM ĐỊNH', 1)
    score = metrics.get('Score_AI_demo', 0)
    dsr = metrics.get('DSR', 0)
    ltv = metrics.get('LTV', 0)
    
    if score >= 0.7 and (np.isnan(dsr) or dsr <= 0.8) and (np.isnan(ltv) or ltv <= 0.8):
        de_xuat = "☑ ĐỀ XUẤT CHO VAY"
        ly_do = "Hồ sơ đáp ứng các chỉ tiêu tài chính, khả năng trả nợ tốt, tài sản bảo đảm đầy đủ."
    elif score >= 0.5:
        de_xuat = "☑ ĐỀ XUẤT CHO VAY CÓ ĐIỀU KIỆN"
        ly_do = "Hồ sơ cần bổ sung thêm tài sản bảo đảm hoặc điều chỉnh điều kiện vay để giảm rủi ro."
    else:
        de_xuat = "☐ KHÔNG ĐỀ XUẤT CHO VAY"
        ly_do = "Hồ sơ không đạt các chỉ tiêu tài chính tối thiểu, rủi ro cao."
    
    p = doc.add_paragraph(de_xuat, style='Heading 3')
    p.style.font.size = Pt(12)
    p = doc.add_paragraph(f"Lý do: {ly_do}")
    p.style.font.size = Pt(11)
    doc.add_paragraph()
    
    # Chữ ký
    p = doc.add_paragraph('_' * 50)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    table_sign = doc.add_table(rows=2, cols=2)
    table_sign.style = 'Table Grid'
    table_sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_sign.autofit = True
    for row in table_sign.rows:
        for cell in row.cells:
            cell.paragraphs[0].style.font.size = Pt(11)
            cell.paragraphs[0].style.font.name = 'Times New Roman'
    
    cells = table_sign.rows[0].cells
    cells[0].text = 'Người thẩm định'
    cells[1].text = 'Phê duyệt'
    
    cells = table_sign.rows[1].cells
    cells[0].text = '(Ký, ghi rõ họ tên)'
    cells[1].text = '(Ký, ghi rõ họ tên)'
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
def export_to_pdf(data: Dict[str, Any], metrics: Dict[str, Any], schedule_df: pd.DataFrame, analysis: str = "") -> str:
    """Xuất báo cáo thẩm định ra LaTeX để tạo PDF"""
    latex_content = r"""
\documentclass[a4paper,12pt]{article}
\usepackage[utf8]{vietnam}
\usepackage{geometry}
\geometry{a4paper, margin=1in}
\usepackage{booktabs}
\usepackage{array}
\usepackage{xcolor}
\usepackage{colortbl}
\usepackage{setspace}
\usepackage[sc]{mathpazo}
\usepackage{noto}
\renewcommand{\familydefault}{\rmdefault}
\begin{document}

\begin{titlepage}
    \centering
    \vspace*{1cm}
    {\LARGE \textbf{NGÂN HÀNG NÔNG NGHIỆP VÀ PHÁT TRIỂN NÔNG THÔN VIỆT NAM}\par}
    \vspace{1cm}
    {\Huge \textbf{BÁO CÁO THẨM ĐỊNH PHƯƠNG ÁN SỬ DỤNG VỐN}\par}
    \vspace{2cm}
    {\large Ngày báo cáo: """ + dt.date.today().strftime("%d/%m/%Y") + r"""\par}
    \vspace{2cm}
    {\large \textit{Hồ sơ khách hàng: """ + data.get('ten_khach_hang', 'Không xác định') + r"""}\par}
\end{titlepage}

\section*{I. Thông tin khách hàng}
\begin{tabular}{|>{\raggedright\arraybackslash}p{5cm}|>{\raggedright\arraybackslash}p{10cm}|}
    \hline
    \textbf{Chỉ tiêu} & \textbf{Giá trị} \\ \hline
    Họ và tên & """ + data.get('ten_khach_hang', '') + r""" \\ \hline
    CMND/CCCD & """ + data.get('cccd', '') + r""" \\ \hline
    Nơi cư trú & """ + data.get('noi_cu_tru', '') + r""" \\ \hline
    Số điện thoại & """ + data.get('so_dien_thoai', '') + r""" \\ \hline
    Mục đích vay & """ + data.get('muc_dich_vay', '') + r""" \\ \hline
\end{tabular}

\vspace{1cm}

\section*{II. Thông tin khoản vay}
\begin{tabular}{|>{\raggedright\arraybackslash}p{5cm}|>{\raggedright\arraybackslash}p{10cm}|}
    \hline
    \textbf{Chỉ tiêu} & \textbf{Giá trị} \\ \hline
    Tổng nhu cầu vốn & """ + format_vnd(data.get('tong_nhu_cau_von', 0)) + r""" VND \\ \hline
    Vốn đối ứng & """ + format_vnd(data.get('von_doi_ung', 0)) + r""" VND \\ \hline
    Số tiền vay & """ + format_vnd(data.get('so_tien_vay', 0)) + r""" VND \\ \hline
    Lãi suất & """ + f"{data.get('lai_suat_nam', 0):.2f}" + r"""\%/năm \\ \hline
    Thời hạn vay & """ + f"{data.get('thoi_gian_vay_thang', 0)}" + r""" tháng \\ \hline
    Thu nhập tháng & """ + format_vnd(data.get('thu_nhap_thang', 0)) + r""" VND \\ \hline
    Giá trị TSĐB & """ + format_vnd(data.get('gia_tri_tsdb', 0)) + r""" VND \\ \hline
\end{tabular}

\vspace{1cm}

\section*{III. Chỉ tiêu tài chính (CADAP)}
\begin{tabular}{|>{\raggedright\arraybackslash}p{5cm}|>{\raggedright\arraybackslash}p{5cm}|>{\raggedright\arraybackslash}p{5cm}|}
    \hline
    \textbf{Chỉ tiêu} & \textbf{Giá trị} & \textbf{Đánh giá} \\ \hline
    PMT (Tiền trả/tháng) & """ + format_vnd(metrics.get('PMT_thang', 0)) + r""" VND & \\ \hline
    DSR (Debt Service Ratio) & """ + (f"{metrics.get('DSR', 0)*100:.1f}\%" if not np.isnan(metrics.get('DSR', 0)) else 'n/a') + r""" & """ + ('✓ Đạt' if (not np.isnan(metrics.get('DSR', 0)) and metrics.get('DSR', 0) <= 0.8) else '✗ Không đạt') + r""" \\ \hline
    LTV (Loan to Value) & """ + (f"{metrics.get('LTV', 0)*100:.1f}\%" if not np.isnan(metrics.get('LTV', 0)) else 'n/a') + r""" & """ + ('✓ Đạt' if (not np.isnan(metrics.get('LTV', 0)) and metrics.get('LTV', 0) <= 0.8) else '✗ Không đạt') + r""" \\ \hline
    E/C (Equity to Capital) & """ + (f"{metrics.get('E_over_C', 0)*100:.1f}\%" if not np.isnan(metrics.get('E_over_C', 0)) else 'n/a') + r""" & """ + ('✓ Đạt' if (not np.isnan(metrics.get('E_over_C', 0)) and metrics.get('E_over_C', 0) >= 0.2) else '✗ Không đạt') + r""" \\ \hline
    CFR (Cash Flow Ratio) & """ + (f"{metrics.get('CFR', 0)*100:.1f}\%" if not np.isnan(metrics.get('CFR', 0)) else 'n/a') + r""" & """ + ('✓ Đạt' if (not np.isnan(metrics.get('CFR', 0)) and metrics.get('CFR', 0) > 0) else '✗ Không đạt') + r""" \\ \hline
    Coverage (Collateral Coverage) & """ + (f"{metrics.get('Coverage', 0)*100:.1f}\%" if not np.isnan(metrics.get('Coverage', 0)) else 'n/a') + r""" & """ + ('✓ Đạt' if (not np.isnan(metrics.get('Coverage', 0)) and metrics.get('Coverage', 0) > 1.2) else '✗ Không đạt') + r""" \\ \hline
    Score tổng hợp & """ + f"{metrics.get('Score_AI_demo', 0):.3f}" + r""" & """ + ('✓ Tốt' if metrics.get('Score_AI_demo', 0) >= 0.7 else ('⚠ Trung bình' if metrics.get('Score_AI_demo', 0) >= 0.5 else '✗ Yếu')) + r""" \\ \hline
\end{tabular}

\vspace{1cm}

\section*{IV. Kế hoạch trả nợ (5 kỳ đầu)}
\begin{tabular}{|c|c|r|r|r|r|}
    \hline
    \textbf{Kỳ} & \textbf{Ngày} & \textbf{Tiền lãi} & \textbf{Tiền gốc} & \textbf{Tổng trả} & \textbf{Dư nợ} \\ \hline
"""
    for i in range(min(5, len(schedule_df))):
        row = schedule_df.iloc[i]
        latex_content += f"    {row['Kỳ']} & {row['Ngày thanh toán']} & {format_vnd(row['Tiền lãi'])} & {format_vnd(row['Tiền gốc'])} & {format_vnd(row['Tổng phải trả'])} & {format_vnd(row['Dư nợ còn lại'])} \\\\ \\hline\n"
    
    latex_content += r"""
\end{tabular}

\begin{center}
(Xem file Excel đính kèm để có đầy đủ """ + str(len(schedule_df)) + r""" kỳ thanh toán)
\end{center}

\vspace{1cm}
"""
    if analysis:
        latex_content += r"""
\section*{V. Phân tích và Kết luận (AI)}
""" + analysis.replace('\n', '\\\\\n') + r"""

\vspace{1cm}
"""
    
    # Ý kiến thẩm định
    score = metrics.get('Score_AI_demo', 0)
    dsr = metrics.get('DSR', 0)
    ltv = metrics.get('LTV', 0)
    
    if score >= 0.7 and (np.isnan(dsr) or dsr <= 0.8) and (np.isnan(ltv) or ltv <= 0.8):
        de_xuat = "☑ ĐỀ XUẤT CHO VAY"
        ly_do = "Hồ sơ đáp ứng các chỉ tiêu tài chính, khả năng trả nợ tốt, tài sản bảo đảm đầy đủ."
    elif score >= 0.5:
        de_xuat = "☑ ĐỀ XUẤT CHO VAY CÓ ĐIỀU KIỆN"
        ly_do = "Hồ sơ cần bổ sung thêm tài sản bảo đảm hoặc điều chỉnh điều kiện vay để giảm rủi ro."
    else:
        de_xuat = "☐ KHÔNG ĐỀ XUẤT CHO VAY"
        ly_do = "Hồ sơ không đạt các chỉ tiêu tài chính tối thiểu, rủi ro cao."
    
    latex_content += r"""
\section*{VI. Ý kiến thẩm định}
\textbf{""" + de_xuat + r"""} \\
Lý do: """ + ly_do + r"""

\vspace{1cm}

\begin{center}
\rule{5cm}{0.4pt}
\end{center}

\begin{tabular}{p{7cm}p{7cm}}
Người thẩm định & Phê duyệt \\
(Ký, ghi rõ họ tên) & (Ký, ghi rõ họ tên) \\
\end{tabular}

\end{document}
"""
    return latex_content
# ========================== UI ==========================
st.title("💼 Thẩm định phương án sử dụng vốn (PASDV)")
st.caption("Upload .docx → Trích xuất → Chỉnh sửa → Tính chỉ tiêu → Kế hoạch trả nợ → Phân tích AI → Xuất Excel/ZIP")
with st.sidebar:
    st.header("⚙️ Cấu hình & Gemini")
    model_name = st.selectbox("Model Gemini", ["gemini-2.0-flash-exp", "gemini-1.5-pro", "gemini-1.5-flash"], index=0)
    api_key = st.text_input("API Key Gemini", type="password", help="Hoặc set GENAI_API_KEY trong secrets.")
    if not api_key:
        api_key = st.secrets.get("GENAI_API_KEY", "") if hasattr(st, "secrets") else ""
    st.markdown("---")

uploaded = st.file_uploader("Tải lên hồ sơ phương án pasdv.docx", type=["docx"], help="Chỉ cần một file .docx")
data = FIELD_DEFAULTS.copy()
if uploaded is not None:
    try:
        data.update(extract_from_docx(uploaded.read()))
        st.success("✅ Đã trích xuất sơ bộ từ file.")
    except Exception as e:
        st.warning(f"⚠️ Không đọc được file DOCX: {e}")
st.markdown("""
<style>
.info-box {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    padding: 20px;
    border-radius: 10px;
    margin-bottom: 20px;
    box-shadow: 0 4px 6px rgba(0,0,0,0.1);
}
.info-box h3 {
    color: white;
    margin: 0;
}
</style>
<div class="info-box">
    <h3>📋 1) Thông tin khách hàng & khoản vay</h3>
</div>
""", unsafe_allow_html=True)
col1, col2, col3 = st.columns(3)
with col1:
    data["ten_khach_hang"] = st.text_input("Họ tên KH", value=data["ten_khach_hang"])
    data["cccd"] = st.text_input("CCCD/CMND", value=data["cccd"])
    data["noi_cu_tru"] = st.text_input("Nơi cư trú", value=data["noi_cu_tru"])
    data["so_dien_thoai"] = st.text_input("Số điện thoại", value=data["so_dien_thoai"])
with col2:
    data["muc_dich_vay"] = st.text_input("Mục đích vay", value=data["muc_dich_vay"])
    data["tong_nhu_cau_von"] = vn_money_input("Tổng nhu cầu vốn (VND)", data["tong_nhu_cau_von"])
    data["von_doi_ung"] = vn_money_input("Vốn đối ứng (VND)", data["von_doi_ung"])
    data["so_tien_vay"] = vn_money_input("Số tiền vay (VND)", data["so_tien_vay"])
with col3:
    data["lai_suat_nam"] = vn_percent_input("Lãi suất (%/năm)", data["lai_suat_nam"])
    data["thoi_gian_vay_thang"] = st.number_input("Thời gian vay (tháng)", value=int(data["thoi_gian_vay_thang"]), min_value=1, max_value=480, step=1)
    data["thu_nhap_thang"] = vn_money_input("Thu nhập tháng (VND)", data["thu_nhap_thang"])
    data["gia_tri_tsdb"] = vn_money_input("Giá trị TSĐB (VND)", data["gia_tri_tsdb"])
col4, col5 = st.columns(2)
with col4:
    data["tong_no_hien_tai"] = vn_money_input("Tổng nợ hiện tại (VND)", data["tong_no_hien_tai"])
with col5:
    data["tong_von_dau_tu"] = vn_money_input("Tổng vốn đầu tư (VND)", data["tong_von_dau_tu"])
    data["loi_nhuan_rong_nam"] = vn_money_input("Lợi nhuận ròng năm (VND)", data["loi_nhuan_rong_nam"])
st.markdown("---")
st.subheader("2) Chỉ tiêu tài chính (CADAP)")
metrics = compute_metrics(data)
if go is not None:
    create_metrics_chart(metrics)
else:
    st.warning("⚠️ Không thể vẽ biểu đồ. Vui lòng cài đặt thư viện Plotly.")
mcol1, mcol2, mcol3, mcol4 = st.columns(4)
with mcol1:
    st.metric("PMT (VND/tháng)", f"{format_vnd(metrics['PMT_thang'])}")
    st.metric("DSR (≤80%)", f"{metrics['DSR']*100:,.1f}%" if not np.isnan(metrics["DSR"]) else "n/a")
with mcol2:
    st.metric("LTV (≤80%)", f"{metrics['LTV']*100:,.1f}%" if not np.isnan(metrics["LTV"]) else "n/a")
    st.metric("E/C (≥20%)", f"{metrics['E_over_C']*100:,.1f}%" if not np.isnan(metrics["E_over_C"]) else "n/a")
with mcol3:
    st.metric("Debt/Income (<4)", f"{metrics['Debt_over_Income']:,.2f}" if not np.isnan(metrics["Debt_over_Income"]) else "n/a")
    st.metric("CFR (>0)", f"{metrics['CFR']*100:,.1f}%" if not np.isnan(metrics["CFR"]) else "n/a")
with mcol4:
    st.metric("Coverage (>120%)", f"{metrics['Coverage']*100:,.1f}%" if not np.isnan(metrics["Coverage"]) else "n/a")
    st.metric("Score demo", f"{metrics['Score_AI_demo']:,.3f}")
ok_flag = "✅" if metrics["Phuong_an_hop_ly"] else "⚠️"
st.info(f"{ok_flag} Tổng nhu cầu vốn {'=' if metrics['Phuong_an_hop_ly'] else '≠'} vốn đối ứng + số tiền vay")
st.markdown("---")
st.markdown("""
<div class="info-box">
    <h3>💰 3) Kế hoạch trả nợ</h3>
</div>
""", unsafe_allow_html=True)
schedule_df = build_amortization(
    principal=data["so_tien_vay"],
    annual_rate_pct=data["lai_suat_nam"],
    months=int(data["thoi_gian_vay_thang"]),
    start_date=dt.date.today()
)
styled_table = style_schedule_table(schedule_df)
st.dataframe(styled_table, use_container_width=True, height=400)
out = io.BytesIO()
with pd.ExcelWriter(out, engine="openpyxl") as writer:
    df_data = pd.DataFrame([data])
    for col in ['tong_nhu_cau_von', 'von_doi_ung', 'so_tien_vay', 'thu_nhap_thang',
                'gia_tri_tsdb', 'tong_no_hien_tai', 'loi_nhuan_rong_nam', 'tong_von_dau_tu']:
        if col in df_data.columns:
            df_data[col] = df_data[col].apply(lambda x: format_vnd(x) if x is not None else None)
    df_metrics = pd.DataFrame([metrics])
    for col in ['PMT_thang']:
        if col in df_metrics.columns:
            df_metrics[col] = df_metrics[col].apply(lambda x: format_vnd(x) if x is not None else None)
    for col in ['DSR', 'LTV', 'E_over_C', 'CFR', 'Coverage', 'ROI']:
        if col in df_metrics.columns:
            df_metrics[col] = df_metrics[col].apply(lambda x: f"{x*100:,.2f}%" if not np.isnan(x) else 'n/a')
    df_data.to_excel(writer, sheet_name="Thong_tin", index=False)
    df_metrics.to_excel(writer, sheet_name="Chi_tieu", index=False)
    schedule_df.to_excel(writer, sheet_name="Ke_hoach", index=False)
out.seek(0)
st.download_button("⬇️ Tải Excel", data=out, file_name="ke_hoach_tra_no.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
# Initialize analysis variable to avoid NameError
analysis = ""
st.subheader("4) Phân tích AI (Gemini)")
if api_key and genai is not None:
    with st.spinner("Đang phân tích..."):
        analysis = gemini_analyze(data, metrics, model_name=model_name, api_key=api_key)
    st.markdown("**Kết luận:**")
    st.write(analysis)
else:
    st.warning("Chưa có API key Gemini. Điền API key ở Sidebar để dùng tính năng này.")
    analysis = ""
# Export Report Section
st.markdown("---")
st.subheader("📄 Xuất Báo cáo")
export_format = st.selectbox("Chọn định dạng báo cáo", ["DOCX", "PDF"], index=0)
if export_format == "DOCX" and Document is None:
    st.info("📄 Cài đặt python-docx để xuất báo cáo DOCX")
elif export_format == "PDF":
    st.info("📄 Báo cáo PDF sẽ được tạo bằng LaTeX (yêu cầu texlive-full và texlive-fonts-extra)")
if st.button("⬇️ Tải Báo cáo"):
    file_name = f"bao_cao_tham_dinh_{data.get('ten_khach_hang', 'khach_hang').replace(' ', '_')}_{dt.date.today().strftime('%Y%m%d')}"
    if export_format == "DOCX" and Document is not None:
        docx_buffer = export_to_docx(data, metrics, schedule_df, analysis=analysis)
        st.download_button(
            "Tải Báo cáo DOCX",
            data=docx_buffer,
            file_name=f"{file_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    elif export_format == "PDF":
        try:
            pdf_content = export_to_pdf(data, metrics, schedule_df, analysis=analysis)
            st.download_button(
                "Tải Báo cáo PDF",
                data=pdf_content.encode('utf-8'),
                file_name=f"{file_name}.tex",
                mime="text/plain"
            )
            st.info("Lưu ý: File .tex cần được biên dịch bằng latexmk với PDFLaTeX để tạo PDF.")
        except Exception as e:
            st.error(f"Lỗi khi tạo LaTeX: {e}")
    else:
        st.error("Không thể xuất báo cáo do thiếu thư viện python-docx.")
st.subheader("5) 💬 Trò chuyện với AI về hồ sơ")
if "chat_messages" not in st.session_state:
    st.session_state.chat_messages = []
for msg in st.session_state.chat_messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])
if prompt := st.chat_input("Hỏi AI về hồ sơ này... (VD: Đánh giá khả năng trả nợ? Rủi ro nào cần lưu ý?)"):
    st.session_state.chat_messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)
    with st.chat_message("assistant"):
        if not api_key:
            response = "⚠️ Vui lòng nhập API Key Gemini ở Sidebar để sử dụng chatbox."
            st.warning(response)
        elif genai is None:
            response = "⚠️ Thư viện google-generativeai chưa được cài đặt."
            st.error(response)
        else:
            try:
                with st.spinner("🤔 AI đang suy nghĩ..."):
                    genai.configure(api_key=api_key)
                    model = genai.GenerativeModel(model_name)
                    context = f"""
Bạn là chuyên viên tín dụng chuyên nghiệp. Dưới đây là thông tin hồ sơ vay:
**Thông tin khách hàng:**
- Họ tên: {data['ten_khach_hang']}
- CCCD: {data['cccd']}
- Địa chỉ: {data['noi_cu_tru']}
- SĐT: {data['so_dien_thoai']}
**Phương án vay:**
- Mục đích: {data['muc_dich_vay']}
- Tổng nhu cầu vốn: {format_vnd(data['tong_nhu_cau_von'])} VND
- Vốn đối ứng: {format_vnd(data['von_doi_ung'])} VND
- Số tiền vay: {format_vnd(data['so_tien_vay'])} VND
- Lãi suất: {data['lai_suat_nam']}%/năm
- Thời hạn: {data['thoi_gian_vay_thang']} tháng
- Thu nhập tháng: {format_vnd(data['thu_nhap_thang'])} VND
- Giá trị TSĐB: {format_vnd(data['gia_tri_tsdb'])} VND
**Chỉ tiêu tài chính:**
- PMT (tiền trả hàng tháng): {format_vnd(metrics['PMT_thang'])} VND
- DSR: {metrics['DSR']*100:.1f}% (chuẩn ≤80%)
- LTV: {metrics['LTV']*100:.1f}% (chuẩn ≤80%)
- E/C: {metrics['E_over_C']*100:.1f}% (chuẩn ≥20%)
- CFR: {metrics['CFR']*100:.1f}% (chuẩn >0%)
- Coverage: {metrics['Coverage']*100:.1f}% (chuẩn >120%)
- Score tổng hợp: {metrics['Score_AI_demo']:.3f}
Hãy trả lời câu hỏi sau dựa trên thông tin trên, sử dụng tiếng Việt chuyên nghiệp nhưng dễ hiểu:
"""
                    full_prompt = context + "\n\nCâu hỏi: " + prompt
                    resp = model.generate_content(full_prompt)
                    response = resp.text if resp.text else "⚠️ Không nhận được phản hồi từ AI."
                    st.markdown(response)
            except Exception as e:
                response = f"❌ Lỗi khi gọi Gemini: {str(e)}"
                st.error(response)
        st.session_state.chat_messages.append({"role": "assistant", "content": response})
col_clear, col_export = st.columns([1, 3])
with col_clear:
    if st.button("🗑️ Xóa chat"):
        st.session_state.chat_messages = []
        st.rerun()
